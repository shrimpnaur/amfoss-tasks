
import os
from dotenv import load_dotenv
import telebot
import requests
import csv
from telebot import types 
from io import StringIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

load_dotenv('bottoken.env')

BOT_TOKEN = os.environ.get('BOT_TOKEN')
GOOGLE_BOOKS_API_KEY = os.environ.get('GOOGLE_BOOKS_API_KEY')

bot = telebot.TeleBot(BOT_TOKEN)

@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, "Heyya :)) , PagePal welcomes you to the vast world of books 📚!! Type /help to see the commands.")

@bot.message_handler(commands=['help'])
def commands_visible(message):
    help_commands = (
        "Hey! 🌟 Here is everything you need to get started with PagePal ;)\n"
        "/start = Helps to get started with PagePal. 🚀\n"
        "/help = Helps navigate you through all my commands. 🗺️\n"
        "/book = Tell me your favorite genre that keeps you on the edge of your seat at 3am, and I will recommend some fiery books to you 🔥📚\n"
        "/preview = I will give you a sneak peek of any book that intrigues you with its preview link!! 🔍📖\n"
        "/list = Your own reading list 📖📋\n"
        "/exit = Goodbye until the next reading session 👋😊\n"
    )
    bot.reply_to(message, help_commands)

@bot.message_handler(commands=['book'])
def book_genre(message):
    bot.reply_to(message, "Please give a genre for me to give you some amazing suggestions! 🌟")
    bot.register_next_step_handler(message, book_recs)

def book_recs(message):
    genre = message.text.strip()

    if not genre:
        bot.reply_to(message, "😕 You didn't provide a genre. Please try again with a valid genre!")
        return

    url = f'https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={GOOGLE_BOOKS_API_KEY}&maxResults=10'
    response = requests.get(url)
    
    # Debug
    print(f"Request URL: {url}")
    print(f"Response Status Code: {response.status_code}")
    print(f"Response Content: {response.text}")

    if response.status_code == 200:
        data = respo
        nse.json()
        if 'items' in data:
            csv_file = StringIO()
            writer = csv.writer(csv_file)
            writer.writerow(['📕 Title', '✍️ Authors', '🏢 Publisher','📅 Published Date', '📝 Description', '🔗 Preview Link'])

            for item in data['items']:
                volume_info = item.get('volumeInfo', {})
                title = volume_info.get('title', 'No title available')
                authors = ', '.join(volume_info.get('authors', ['Unknown author']))
                publisher = volume_info.get('publisher', 'No publisher available')
                published_date = item['volumeInfo'].get('publishedDate', 'No published date available')
                description = volume_info.get('description', 'No description available')
                preview_link = volume_info.get('previewLink', 'No preview link available')

                writer.writerow([title, authors, publisher, published_date, description, preview_link])

            csv_file.seek(0)              
            bot.send_document(message.chat.id, csv_file, caption=f"📂 {genre.capitalize()} Books Recommendations just for you ;)")
        else:
            bot.reply_to(message, "😕 Sorry, I couldn't find books in this genre. Try something different!")
    else:
        bot.reply_to(message, f"⚠️ Something went wrong (status code: {response.status_code}). Please try again later.")

@bot.message_handler(commands=['preview'])
def book_preview(message):
    bot.reply_to(message, "Which book would you like to know about? 🤔 Give me the name, and I will provide you with the preview link! 🔗")
    bot.register_next_step_handler(message, preview_link)

def preview_link(message):
    book_name = message.text

    if not book_name:
        bot.reply_to(message, "😕 Oh no! Could you please try again with a valid book name!")
        return

    url = f'https://www.googleapis.com/books/v1/volumes?q={book_name}&key={GOOGLE_BOOKS_API_KEY}'
    response = requests.get(url)

    if response.status_code == 200:
        data = response.json()
        if 'items' in data and len(data['items']) > 0:
            volume_info = data['items'][0].get('volumeInfo', {})
            title = volume_info.get('title', 'No title available')
            preview_link = volume_info.get('previewLink', None)

            if preview_link:
                bot.reply_to(message, f"📖 Here is the preview link for '{title}': {preview_link}")
            else:
                bot.reply_to(message, f"😕 Sorry, no preview link available for '{title}'.")
        else:
            bot.reply_to(message, "😕 Sorry, I couldn't find any information on that book. Please try again with a different name.")
    else:
        bot.reply_to(message, f"⚠️ Something went wrong (status code: {response.status_code}). Please try again later.\n\nResponse Content: {response.text}")

user_states = {}
@bot.message_handler(commands=['list'])
def list_book(message):
     user_states[message.chat.id] = 'waiting_for_reading_list'
     bot.reply_to(message, "Okay, time for working with the reading list! 📖🗒️ Please enter the /reading_list command to add, delete, or view the list.")

readinglist_file = "reading_list.docx"

def make_read_doc():
    if not os.path.exists(readinglist_file):
        doc = Document()

        title = doc.add_heading('My Reading List', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        
        run = title.runs[0]
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        doc.add_paragraph(' ')
        doc.save(readinglist_file)

@bot.message_handler(commands=['reading_list'])
def reading_list(message):
    if user_states.get(message.chat.id) == 'waiting_for_reading_list':
        make_read_doc()
        markup = types.InlineKeyboardMarkup()

        butn_add = types.InlineKeyboardButton('Add a book 📚', callback_data='add')
        butn_delete = types.InlineKeyboardButton('Delete a book ❌', callback_data='delete')
        butn_view = types.InlineKeyboardButton('View Reading List 📋', callback_data='view')
        markup.add(butn_add, butn_delete, butn_view)
        bot.send_message(message.chat.id, "So, what do you plan on doing with your list? 📋", reply_markup=markup)
        user_states[message.chat.id] = None
    else:
        bot.reply_to(message, "Oops, there might be a spelling mistake! 😅 Please use the /list command first.")

@bot.callback_query_handler(func=lambda call: True)
def handle_list_action(call):
    action = call.data

    if action == 'add':
        bot.send_message(call.message.chat.id, "Please type the name of the book you want to add 📚:")
        bot.register_next_step_handler_by_chat_id(call.message.chat.id, add_process)
    elif action == 'delete':
        bot.send_message(call.message.chat.id, "Please type the name of the book you want to delete ❌:")
        bot.register_next_step_handler_by_chat_id(call.message.chat.id, del_process)
    elif action == 'view':
        view_reading_list(call.message)
    else:
        bot.reply_to(call.message, "❌ Invalid action. Please try again.")
        reading_list(call.message)

def add_process(message):
    book_name = message.text.strip()

    if not book_name:
        bot.reply_to(message, "❌ Invalid book name. Please try again.")
        return
    
    url = f'https://www.googleapis.com/books/v1/volumes?q={book_name}&key={GOOGLE_BOOKS_API_KEY}'
    response = requests.get(url)

    if response.status_code == 200:
        data = response.json()
        if 'items' in data and len(data['items']) > 0:
            volume_info = data['items'][0].get('volumeInfo', {})
            title = volume_info.get('title', book_name)
            preview_link = volume_info.get('previewLink', 'No preview link available')

            # Checking for redundant books
            make_read_doc()
            doc = Document(readinglist_file)
            paragraphs = doc.paragraphs
            found = False
            for para in paragraphs:
                if para.text.startswith("📕") and title.lower() in para.text.lower():
                    bot.reply_to(message, f"'{title}' is already in your reading list! 📚")
                    found = True
                    break

            if not found:
                doc.add_paragraph(f"📕 {title}")
                doc.add_paragraph(f"🔗 {preview_link}")
                doc.add_paragraph(" ") 
            
            doc.paragraphs[-3].style.font.size = Pt(14)
            doc.paragraphs[-3].style.font.bold = True
            doc.paragraphs[-3].style.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)


            
            doc.save(readinglist_file)

            bot.reply_to(message, f"✅ Added '{title}' to your reading list!")
        else:
            bot.reply_to(message, "😕 Sorry, I couldn't find that book. Please try again with a different name.")
    else:
        bot.reply_to(message, f"⚠️ Something went wrong (status code: {response.status_code}). Please try again later.")

def del_process(message):
    book_name = message.text.strip()
    if not book_name:
        bot.reply_to(message, "❌ Invalid book name. Please try again.")
        return
    
    doc = Document(readinglist_file)
    paragraphs = doc.paragraphs
    paragraphs_to_remove = []
    found = False


    for i in range(len(paragraphs)):
        if paragraphs[i].text.startswith("📕") and book_name.lower() in paragraphs[i].text.lower():
            found = True
            paragraphs_to_remove = [i, i+1, i+2]
            break
      
    if found:
        for i in reversed(paragraphs_to_remove):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
        doc.save(readinglist_file)
        bot.reply_to(message, f"🗑️ Deleted '{book_name}' from your reading list!")
    else:
        bot.reply_to(message, f"😕 Sorry, '{book_name}' is not in your reading list.")


    
         
def view_reading_list(message):
     with open(readinglist_file, "rb") as doc:
        bot.send_document(message.chat.id, doc, caption="📄 Here is your amazing reading list ;))")

@bot.message_handler(commands=['exit'])
def exit_bot(message):
    bot.reply_to(message, "Now I take my leave 👋📚 Keep reading buddy until the next book adventure!")

bot.infinity_polling()

















